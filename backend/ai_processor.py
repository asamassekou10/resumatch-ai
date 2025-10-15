import PyPDF2
import docx
import re
import hashlib
import pickle
import os
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import spacy
from io import BytesIO
from functools import lru_cache
from typing import Dict, List, Tuple, Optional
from datetime import datetime
from errors import AIProcessingError, FileProcessingError
import logging

logger = logging.getLogger(__name__)

class AIProcessor:
    """Enhanced AI processor with caching and error handling"""
    
    def __init__(self, cache_dir='cache'):
        self.cache_dir = cache_dir
        self.nlp = None
        self.vectorizer = None
        self._initialize_models()
        self._ensure_cache_dir()
    
    def _initialize_models(self):
        """Initialize AI models once at startup"""
        try:
            # Load spaCy model
            self.nlp = spacy.load("en_core_web_sm")
            logger.info("spaCy model loaded successfully")
        except OSError:
            logger.warning("spaCy model not found. Run: python -m spacy download en_core_web_sm")
            self.nlp = None
        
        # Initialize TF-IDF vectorizer
        self.vectorizer = TfidfVectorizer(
            stop_words='english',
            ngram_range=(1, 3),
            max_features=1000,
            min_df=1
        )
        logger.info("TF-IDF vectorizer initialized")
    
    def _ensure_cache_dir(self):
        """Ensure cache directory exists"""
        if not os.path.exists(self.cache_dir):
            os.makedirs(self.cache_dir)
    
    def _generate_cache_key(self, resume_text: str, job_description: str) -> str:
        """Generate cache key for resume and job description combination"""
        combined = f"{resume_text[:1000]}{job_description[:1000]}"
        return hashlib.md5(combined.encode()).hexdigest()
    
    def _load_from_cache(self, cache_key: str) -> Optional[Dict]:
        """Load analysis result from cache"""
        cache_file = os.path.join(self.cache_dir, f"{cache_key}.pkl")
        if os.path.exists(cache_file):
            try:
                with open(cache_file, 'rb') as f:
                    return pickle.load(f)
            except Exception as e:
                logger.warning(f"Failed to load from cache: {e}")
        return None
    
    def _save_to_cache(self, cache_key: str, result: Dict):
        """Save analysis result to cache"""
        cache_file = os.path.join(self.cache_dir, f"{cache_key}.pkl")
        try:
            with open(cache_file, 'wb') as f:
                pickle.dump(result, f)
        except Exception as e:
            logger.warning(f"Failed to save to cache: {e}")
    
    def extract_text_from_pdf(self, file_stream: BytesIO) -> str:
        """Extract text from PDF file with enhanced error handling"""
        try:
            pdf_reader = PyPDF2.PdfReader(file_stream)
            text = ""
            
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
                except Exception as e:
                    logger.warning(f"Failed to extract text from page {page_num}: {e}")
                    continue
            
            if not text.strip():
                raise FileProcessingError("PDF appears to be empty or contains only images")
            
            return text.strip()
            
        except Exception as e:
            raise FileProcessingError(f"PDF parsing failed: {str(e)}")
    
    def extract_text_from_docx(self, file_stream: BytesIO) -> str:
        """Extract text from DOCX file with enhanced error handling"""
        try:
            doc = docx.Document(file_stream)
            text_parts = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text.strip())
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text.strip())
            
            text = "\n".join(text_parts)
            
            if not text.strip():
                raise FileProcessingError("DOCX appears to be empty")
            
            return text.strip()
            
        except Exception as e:
            raise FileProcessingError(f"DOCX parsing failed: {str(e)}")
    
    def extract_text_from_file(self, file) -> str:
        """Extract text from uploaded file with validation"""
        filename = file.filename.lower()
        file_stream = BytesIO(file.read())
        
        try:
            if filename.endswith('.pdf'):
                return self.extract_text_from_pdf(file_stream)
            elif filename.endswith('.docx'):
                return self.extract_text_from_docx(file_stream)
            elif filename.endswith('.txt'):
                content = file_stream.read().decode('utf-8')
                if not content.strip():
                    raise FileProcessingError("TXT file appears to be empty")
                return content.strip()
            else:
                raise FileProcessingError("Unsupported file format. Please upload PDF, DOCX, or TXT")
                
        except UnicodeDecodeError:
            raise FileProcessingError("File contains invalid characters. Please ensure it's a valid text file.")
        except Exception as e:
            if isinstance(e, (FileProcessingError, AIProcessingError)):
                raise
            raise FileProcessingError(f"File processing failed: {str(e)}")
    
    @lru_cache(maxsize=128)
    def extract_keywords_tfidf(self, text: str, top_n: int = 20) -> List[str]:
        """Extract keywords using TF-IDF with caching"""
        try:
            # Clean and preprocess text
            cleaned_text = self._preprocess_text(text)
            
            if not cleaned_text:
                return []
            
            # Use TF-IDF to find important keywords
            vectorizer = TfidfVectorizer(
                max_features=top_n * 2,  # Get more features for better filtering
                stop_words='english',
                ngram_range=(1, 3),
                min_df=1,
                max_df=0.8  # Ignore terms that appear in more than 80% of documents
            )
            
            tfidf_matrix = vectorizer.fit_transform([cleaned_text])
            feature_names = vectorizer.get_feature_names_out()
            
            # Get scores for each keyword
            scores = tfidf_matrix.toarray()[0]
            keyword_scores = list(zip(feature_names, scores))
            keyword_scores.sort(key=lambda x: x[1], reverse=True)
            
            # Filter out very short keywords and return top results
            filtered_keywords = [
                kw for kw, score in keyword_scores 
                if score > 0 and len(kw) > 2
            ]
            
            return filtered_keywords[:top_n]
            
        except Exception as e:
            logger.warning(f"TF-IDF keyword extraction failed: {e}")
            # Fallback to simple word frequency
            return self._fallback_keyword_extraction(text, top_n)
    
    def _fallback_keyword_extraction(self, text: str, top_n: int) -> List[str]:
        """Fallback keyword extraction using simple word frequency"""
        words = text.lower().split()
        # Filter out common words and short words
        stop_words = {'the', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for', 'of', 'with', 'by', 'a', 'an'}
        word_freq = {}
        
        for word in words:
            clean_word = re.sub(r'[^a-zA-Z]', '', word)
            if len(clean_word) > 3 and clean_word not in stop_words:
                word_freq[clean_word] = word_freq.get(clean_word, 0) + 1
        
        return sorted(word_freq.keys(), key=lambda x: word_freq[x], reverse=True)[:top_n]
    
    def extract_skills_with_ner(self, text: str) -> List[str]:
        """Extract skills using Named Entity Recognition with spaCy"""
        if not self.nlp:
            return self._extract_skills_fallback(text)
        
        try:
            doc = self.nlp(text)
            skills = set()
            
            # Common technical skills patterns
            skill_patterns = [
                'python', 'java', 'javascript', 'typescript', 'react', 'vue', 'angular', 'node',
                'sql', 'mongodb', 'postgresql', 'mysql', 'redis', 'elasticsearch',
                'aws', 'azure', 'gcp', 'docker', 'kubernetes', 'terraform',
                'machine learning', 'data analysis', 'api', 'rest', 'graphql',
                'agile', 'scrum', 'git', 'ci/cd', 'testing', 'flask', 'django',
                'html', 'css', 'bootstrap', 'tailwind', 'express', 'redux',
                'linux', 'windows', 'macos', 'bash', 'powershell'
            ]
            
            text_lower = text.lower()
            for skill in skill_patterns:
                if skill in text_lower:
                    skills.add(skill)
            
            # Extract entities that might be skills
            for ent in doc.ents:
                if ent.label_ in ['ORG', 'PRODUCT', 'GPE', 'TECHNOLOGY']:
                    if 3 <= len(ent.text) <= 30:
                        skills.add(ent.text.lower())
            
            # Extract technical terms using POS tagging
            for token in doc:
                if token.pos_ in ['NOUN', 'PROPN'] and len(token.text) > 3:
                    # Check if it looks like a technical term
                    if any(char.isupper() for char in token.text) or \
                       token.text.lower() in skill_patterns:
                        skills.add(token.text.lower())
            
            return list(skills)[:20]  # Limit to top 20 skills
            
        except Exception as e:
            logger.warning(f"NER skill extraction failed: {e}")
            return self._extract_skills_fallback(text)
    
    def _extract_skills_fallback(self, text: str) -> List[str]:
        """Fallback skill extraction without spaCy"""
        # Simple pattern-based skill extraction
        skills = set()
        
        # Technical skills patterns
        tech_patterns = [
            r'\b(python|java|javascript|typescript|react|vue|angular|node)\b',
            r'\b(sql|mongodb|postgresql|mysql|redis)\b',
            r'\b(aws|azure|gcp|docker|kubernetes)\b',
            r'\b(html|css|bootstrap|tailwind)\b',
            r'\b(git|github|gitlab)\b',
            r'\b(linux|windows|macos)\b'
        ]
        
        text_lower = text.lower()
        for pattern in tech_patterns:
            matches = re.findall(pattern, text_lower)
            skills.update(matches)
        
        return list(skills)[:15]
    
    def calculate_match_score(self, resume_text: str, job_description: str) -> float:
        """Calculate similarity score between resume and job description"""
        try:
            # Preprocess texts
            resume_clean = self._preprocess_text(resume_text)
            job_clean = self._preprocess_text(job_description)
            
            if not resume_clean or not job_clean:
                return 0.0
            
            # Use TF-IDF for similarity calculation
            vectorizer = TfidfVectorizer(
                stop_words='english',
                ngram_range=(1, 2),
                max_features=500
            )
            
            tfidf_matrix = vectorizer.fit_transform([resume_clean, job_clean])
            similarity = cosine_similarity(tfidf_matrix[0:1], tfidf_matrix[1:2])[0][0]
            
            # Convert to percentage and round
            return round(similarity * 100, 2)
            
        except Exception as e:
            logger.warning(f"Match score calculation failed: {e}")
            # Fallback: simple word overlap
            return self._calculate_overlap_score(resume_text, job_description)
    
    def _calculate_overlap_score(self, resume_text: str, job_description: str) -> float:
        """Fallback overlap score calculation"""
        try:
            resume_words = set(re.findall(r'\b\w+\b', resume_text.lower()))
            job_words = set(re.findall(r'\b\w+\b', job_description.lower()))
            
            # Remove common stop words
            stop_words = {'the', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for', 'of', 'with', 'by', 'a', 'an', 'is', 'are', 'was', 'were', 'be', 'been', 'being', 'have', 'has', 'had', 'do', 'does', 'did', 'will', 'would', 'could', 'should', 'may', 'might', 'must', 'can'}
            resume_words -= stop_words
            job_words -= stop_words
            
            if not job_words:
                return 0.0
            
            overlap = len(resume_words.intersection(job_words))
            return round((overlap / len(job_words)) * 100, 2)
            
        except Exception:
            return 0.0
    
    def _preprocess_text(self, text: str) -> str:
        """Preprocess text for analysis"""
        if not text:
            return ""
        
        # Convert to lowercase
        text = text.lower()
        
        # Remove extra whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove special characters but keep letters, numbers, and spaces
        text = re.sub(r'[^a-zA-Z0-9\s+#]', ' ', text)
        
        return text.strip()
    
    def generate_suggestions(self, keywords_missing: List[str], match_score: float, 
                           keywords_found: List[str]) -> str:
        """Generate actionable suggestions for the user"""
        suggestions = []
        
        # Score-based suggestions
        if match_score < 30:
            suggestions.append("Your resume shows very limited alignment with this job description. Consider a significant revision to highlight relevant experience and skills.")
        elif match_score < 50:
            suggestions.append("Your resume needs improvement to better match this job. Focus on incorporating relevant keywords and experiences.")
        elif match_score < 70:
            suggestions.append("Your resume is moderately aligned. Focus on incorporating the missing keywords where they genuinely reflect your experience.")
        else:
            suggestions.append("Strong alignment! Your resume matches well with this job description.")
        
        # Keywords-based suggestions
        if len(keywords_missing) > 15:
            suggestions.append(f"Consider adding {len(keywords_missing)} key skills from the job description. Focus on: {', '.join(keywords_missing[:5])}.")
        elif len(keywords_missing) > 5:
            suggestions.append(f"Add these important keywords where relevant: {', '.join(keywords_missing[:5])}.")
        elif len(keywords_missing) > 0:
            suggestions.append(f"Minor improvements: consider adding {', '.join(keywords_missing[:3])}.")
        else:
            suggestions.append("Excellent! You've covered all major keywords from the job description.")
        
        # General suggestions
        suggestions.append("Use action verbs and quantify your achievements with specific numbers and results.")
        suggestions.append("Tailor your professional summary to match the job requirements and company culture.")
        suggestions.append("Ensure your experience section highlights the most relevant projects and responsibilities.")
        
        # Format suggestions
        if match_score < 60:
            suggestions.append("Consider restructuring your resume to lead with your most relevant experience first.")
        
        return " ".join(suggestions)
    
    def process_resume_analysis(self, resume_file, job_description: str) -> Dict:
        """Main function to process resume analysis with caching"""
        try:
            # Extract text from resume
            resume_text = self.extract_text_from_file(resume_file)
            
            if len(resume_text) < 50:
                raise FileProcessingError("Resume text is too short. Please ensure the file contains sufficient content.")
            
            # Check cache first
            cache_key = self._generate_cache_key(resume_text, job_description)
            cached_result = self._load_from_cache(cache_key)
            
            if cached_result:
                logger.info("Using cached analysis result")
                return cached_result
            
            # Extract keywords from both documents
            resume_keywords = set(self.extract_keywords_tfidf(resume_text, top_n=30))
            resume_skills = set(self.extract_skills_with_ner(resume_text))
            resume_all = resume_keywords.union(resume_skills)
            
            job_keywords = set(self.extract_keywords_tfidf(job_description, top_n=30))
            job_skills = set(self.extract_skills_with_ner(job_description))
            job_all = job_keywords.union(job_skills)
            
            # Find matches and gaps
            keywords_found = list(resume_all.intersection(job_all))
            keywords_missing = list(job_all - resume_all)
            
            # Calculate match score
            match_score = self.calculate_match_score(resume_text, job_description)
            
            # Generate suggestions
            suggestions = self.generate_suggestions(keywords_missing, match_score, keywords_found)
            
            result = {
                'match_score': match_score,
                'keywords_found': keywords_found[:20],  # Limit to top 20
                'keywords_missing': keywords_missing[:20],  # Limit to top 20
                'suggestions': suggestions,
                'resume_text': resume_text[:1000],  # Store excerpt for reference
                'analysis_timestamp': datetime.utcnow().isoformat()
            }
            
            # Cache the result
            self._save_to_cache(cache_key, result)
            
            logger.info(f"Analysis completed successfully. Match score: {match_score}%")
            return result
            
        except Exception as e:
            if isinstance(e, (FileProcessingError, AIProcessingError)):
                raise
            logger.error(f"Resume analysis failed: {e}")
            raise AIProcessingError(f"Resume analysis failed: {str(e)}")

# Create global instance
ai_processor = AIProcessor()

# Backward compatibility function
def process_resume_analysis(resume_file, job_description):
    """Backward compatibility function"""
    return ai_processor.process_resume_analysis(resume_file, job_description)
